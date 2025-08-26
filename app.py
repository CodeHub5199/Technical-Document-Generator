import streamlit as st
from crewai import Agent, Task, Crew, LLM
from docx import Document
from docx.shared import Pt
import os
from dotenv import load_dotenv
import io
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter

load_dotenv()

# ----------------------------
# File Processing
# ----------------------------
def process_uploaded_file(uploaded_file):
    """Process uploaded file with chunking for large files"""
    try:
        text = uploaded_file.read().decode('utf-8')
        
        # Chunk only if exceeds Groq's context window (8192 tokens)
        if len(text) > 3000:  # Conservative estimate (~3000 chars ≈ 2000 tokens)
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=2000,  # Safe chunk size in characters
                chunk_overlap=200,
                length_function=len
            )
            return splitter.split_text(text)
        return [text]
    except Exception as e:
        st.error(f"File processing error: {str(e)}")
        return []

# ----------------------------
# Improved Document Generator (Modularized)
# ----------------------------
class DocxGenerator:
    def __init__(self):
        self.doc = Document()
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri (Body)'
        font.size = Pt(12)
        self.headings = []
    
    def _add_heading_to_toc(self, level, text):
        """Efficiently cache headings for TOC"""
        if level > 1:  # Skip document title
            self.headings.append((level, text))
    
    def _process_text_segments(self, paragraph, text):
        """Optimized text formatting with regex"""
        for segment in re.split(r'(\*\*.+?\*\*)', text):
            if segment.startswith('**') and segment.endswith('**'):
                run = paragraph.add_run(segment[2:-2])
                run.bold = True
            else:
                paragraph.add_run(segment)

    def markdown_to_word(self, markdown_text):
        """Optimized markdown parser using single-pass processing"""
        current_solution_heading = None
        
        for section in re.split(r'(^#+.+$)', markdown_text, flags=re.MULTILINE):
            if not section.strip():
                continue
                
            if section.startswith('#'):
                heading_level = section.count('#')
                heading_text = section.replace('#', '').strip()
                self._add_heading_to_toc(heading_level, heading_text)
                
                # Smart heading nesting
                adjusted_level = min(heading_level, 2)
                if "Solution" in heading_text:
                    current_solution_heading = heading_text
                elif "How It Works" in heading_text and current_solution_heading:
                    adjusted_level = 3
                
                self.doc.add_heading(heading_text, level=adjusted_level)
            else:
                for line in section.split('\n'):
                    if not line.strip():
                        continue
                        
                    if line.strip().startswith('- '):
                        p = self.doc.add_paragraph(style='List Bullet')
                        self._process_text_segments(p, line[2:])
                    elif re.match(r'^\d+\. ', line.strip()):
                        p = self.doc.add_paragraph(style='List Number')
                        self._process_text_segments(p, re.sub(r'^\d+\. ', '', line))
                    else:
                        p = self.doc.add_paragraph()
                        self._process_text_segments(p, line)

    def save_to_buffer(self):
        """Memory-efficient buffer handling"""
        doc_buffer = io.BytesIO()
        self.doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer

# ----------------------------
# Agent Definitions (Optimized)
# ----------------------------
def create_agents():
    GROQ_API_KEY = os.getenv('GROQ_API_KEY')
    llm = LLM(model="groq/llama-3.3-70b-versatile", temperature=0, max_tokens=3000)
    
    code_analyst = Agent(
        role='Code Analyst',
        goal='Analyze original vs modified code',
        backstory='Expert in differential code analysis',
        llm=llm,
        max_iter=3
    )
    
    document_architect = Agent(
        role='Document Architect',
        goal='Generate technical documentation',
        backstory='Technical writer specializing in code documentation',
        llm=llm,
        max_iter=2
    )
    
    return code_analyst, document_architect

# ----------------------------
# Enhanced File Processing
# ----------------------------
def extract_relevant_code(full_code, changed_code):
    """Extract contextually relevant portions from original code"""
    try:
        # Find import statements and class/function definitions
        pattern = r'^(import .+|from .+ import .+|def \w+\(|class \w+:)'
        relevant_lines = []
        
        for line in full_code.split('\n'):
            if re.search(pattern, line):
                relevant_lines.append(line)
            # Keep 5 lines before/after matches for context
            elif relevant_lines and len(relevant_lines[-1].split('\n')) < 5:
                relevant_lines.append(line)
        
        # Add changed code references if found in original
        for changed_line in changed_code.split('\n')[:10]:  # First 10 lines of changed code
            if changed_line.strip() and changed_line in full_code:
                idx = full_code.index(changed_line)
                context = full_code[max(0, idx-200):min(len(full_code), idx+200)]
                relevant_lines.append(f"\n# Context for changes:\n{context}")
        
        return '\n'.join(relevant_lines[:5000])  # Safe limit
    except:
        return full_code[:5000]  # Fallback

# ----------------------------
# Streamlit UI
# ----------------------------
st.set_page_config(page_title="Design Document Generator", layout="wide")
st.title("Design Document Generator")

st.subheader("User Story Details")
user_story_name = st.text_input("User Story Name")
user_story_description = st.text_area("User Story Description")

# File upload for original code
original_file = st.file_uploader(
    "Upload Original Code File",
    type=['py', 'txt', 'js', 'java', 'c', 'cpp', 'go'],
    help="Upload the original version of your code file"
)

# Textarea for changed code
code_snippet = st.text_area(
    "Paste Your Modified Code",
    height=300,
    help="Paste only the changed/new code sections"
)

# Additional Prompt Section (Always visible)
st.subheader("Additional Instructions & Context")
additional_prompt = st.text_area(
    "Provide additional context or specific instructions for the analysis",
    height=150,
    help="Add any extra information, specific focus areas, or domain context that will help the agents understand your code better",
    placeholder="Example: This change is part of a performance optimization effort. Please focus on memory usage implications..."
)

if st.button("Explain Code Changes"):
    with st.spinner("Analyzing changes..."):
        try:
            # Get and process original code
            original_code = ""
            if original_file:
                original_code = original_file.read().decode('utf-8')
            original_code = extract_relevant_code(original_code, code_snippet) if original_code else ""

            # Initialize agents
            code_analyst, document_architect = create_agents()

            # Build task description with additional prompt
            task_description = f"""Analyze these code changes:
            
            Original Code (Key Sections):
            {original_code or 'No original code provided'}
            
            Modified Code:
            {code_snippet or 'No changes detected'}"""

            # Add additional prompt if provided
            if additional_prompt.strip():
                task_description += f"""
                
            Additional Context & Instructions:
            {additional_prompt}"""

            task_description += """
            
            Provide technical explanation in this exact format:
            
            ## Solution
            [Overall what changed]
            
            ### How It Works
            [Technical details with code references]
            
            ### Impacts
            [Potential effects on system]"""

            # Create analysis task
            analysis_task = Task(
                description=task_description,
                agent=code_analyst,
                expected_output="Structured technical explanation in specified format",
            )

            # Execute workflow
            crew = Crew(
                agents=[code_analyst, document_architect],
                tasks=[analysis_task],
                verbose=True
            )
            
            # Get result and generate Word doc
            result = str(crew.kickoff())
            
            doc_generator = DocxGenerator()
            
            # 1. Add Title
            doc_generator.doc.add_heading('User Story Name', level=2)
            doc_generator.doc.add_paragraph(user_story_name)
            
            # 2. Add User Story Description
            if user_story_description:
                doc_generator.doc.add_heading('User Story Description', level=2)
                doc_generator.doc.add_paragraph(user_story_description)
            
            # 3. Add Additional Context section if provided
            if additional_prompt.strip():
                doc_generator.doc.add_heading('Additional Context & Instructions', level=2)
                doc_generator.doc.add_paragraph(additional_prompt)
            
            # 4. Add Code Analysis (directly from result)
            doc_generator.markdown_to_word(result)  # Convert to Word formatting
            
            # 5. Save and create download button
            doc_buffer = doc_generator.save_to_buffer()
            
            st.subheader("Analysis Complete")
            st.download_button(
                label="Download Word Document",
                data=doc_buffer,
                file_name=f"{(user_story_name or 'code_analysis').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Optional: Show preview
            with st.expander("Preview Analysis"):
                st.markdown(result)
                
        except Exception as e:
            st.error(f"Analysis failed: {str(e)}")